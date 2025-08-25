import streamlit as st
import os
import io
import tempfile
import requests
import time
import numpy as np
import re
from urllib.parse import urljoin
import pandas as pd
import difflib  # fuzzy matching nama kolom

# Mistral 0.4.2 (legacy client)
from mistralai.client import MistralClient

import google.generativeai as genai
from PIL import Image
from PyPDF2 import PdfReader

# --- Tambahan import untuk dashboard ---
import plotly.express as px
import plotly.graph_objects as go
from collections import Counter
import json
import string
import subprocess

# --- DOCX import with friendly error ---
try:
    # modul bernama "docx" disediakan oleh paket "python-docx"
    from docx import Document  # pip install python-docx
except Exception as e:
    st.error(
        "Paket yang benar untuk .docx adalah **python-docx**. "
        "Sepertinya paket **docx** (yang salah) terpasang dan menyebabkan konflik.\n\n"
        "Perbaikan cepat:\n"
        "1) aktifkan venv\n"
        "2) `pip uninstall -y docx`\n"
        "3) `pip install --upgrade python-docx`\n\n"
        f"Detail error: {e}"
    )
    raise

# WordCloud opsional (fallback otomatis kalau belum terpasang)
try:
    from wordcloud import WordCloud
    HAS_WORDCLOUD = True
except Exception:
    HAS_WORDCLOUD = False


# --------------------- Page config ---------------------
st.set_page_config(page_title="Document Intelligence Agent", layout="wide")
st.title("Document Intelligence Agent")
st.markdown("Upload documents or URL to extract information and ask questions")


# --------------------- Sidebar: API Keys ----------------
with st.sidebar:
    st.header("API Configuration")
    mistral_api_key = st.text_input("Mistral AI API Key (legacy 0.4.2)", type="password")
    google_api_key = st.text_input("Google API Key (Gemini)", type="password")
    model_preference = st.selectbox(
        "Model preference",
        ["Auto (Pro→Flash)", "Flash only", "Pro only"],
        index=0,
        help="Fallback otomatis ke Flash saat Pro kena rate limit/kuota."
    )
    answer_language = st.selectbox(
        "Answer language",
        ["Bahasa Indonesia", "English"],
        index=0,
        help="Bahasa jawaban untuk fitur Q&A."
    )

    st.markdown("---")
    st.markdown("How To Get API Key Tutorials")
    st.markdown(
        """
- **Mistral AI API Key** — [YouTube Tutorial](https://youtu.be/NUCcUFwfhlA?si=iLrFFxVtcFUp657C)  
- **Google API Key (Gemini)** — [YouTube Tutorial](https://youtu.be/IHj7wF-8ry8?si=VKvhMM3pMeKwkXAv)
        """
    )


# Disimpan global agar helper bisa akses
MODEL_PREFERENCE = model_preference
ANSWER_LANGUAGE = answer_language

# MistralClient (legacy) — tidak untuk OCR di 0.4.2
mistral_client = None
if mistral_api_key:
    try:
        mistral_client = MistralClient(api_key=mistral_api_key)
        st.success("✅ Mistral API connected (legacy client 0.4.2)")
    except Exception as e:
        st.error(f"Failed to initialize Mistral client: {e}")

# Gemini untuk OCR & QnA
if google_api_key:
    try:
        genai.configure(api_key=google_api_key)
        st.success("✅ Google API connected")
    except Exception as e:
        st.error(f"Failed to initialize Google API: {e}")

# --------------------- Helpers (Gemini OCR) -------------------------
def _is_quota_error(err: Exception) -> bool:
    msg = str(err).lower()
    return "429" in msg or "quota" in msg or "rate limit" in msg or "exceeded" in msg

def _get_model_candidates() -> list:
    if MODEL_PREFERENCE == "Flash only":
        return ["gemini-1.5-flash"]
    if MODEL_PREFERENCE == "Pro only":
        return ["gemini-1.5-pro"]
    return ["gemini-1.5-pro", "gemini-1.5-flash"]

def _generate_with_fallback(parts_or_prompt):
    last_error = None
    candidates = _get_model_candidates()
    for model_name in candidates:
        try:
            model = genai.GenerativeModel(model_name=model_name)
            resp = model.generate_content(parts_or_prompt)
            text = getattr(resp, "text", "").strip()
            if text:
                if model_name != candidates[0]:
                    st.info(f"Using fallback model: {model_name}")
                return text
        except Exception as e:
            last_error = e
            if _is_quota_error(e):
                continue
            else:
                return f"Error generating response: {e}"
    if last_error and _is_quota_error(last_error):
        time.sleep(30)
        try:
            model = genai.GenerativeModel(model_name=candidates[-1])
            resp = model.generate_content(parts_or_prompt)
            return getattr(resp, "text", "").strip()
        except Exception as e2:
            return f"Error generating response: {e2}"
    return f"Error generating response: {last_error}"

def _truncate_context(text: str, max_chars: int = 20000) -> str:
    if not text:
        return text
    if len(text) <= max_chars:
        return text
    return text[:max_chars]

# --------------------- Text chunking & Embeddings (RAG) ---------------------
EMBEDDING_MODEL = "models/text-embedding-004"

def _chunk_text(text: str, chunk_size: int = 1800, overlap: int = 200) -> list:
    if not text:
        return []
    chunks = []
    start = 0
    end = max(chunk_size, 1)
    text_len = len(text)
    while start < text_len:
        chunk = text[start:end]
        chunks.append(chunk)
        if end >= text_len:
            break
        start = max(end - overlap, 0)
        end = min(start + chunk_size, text_len)
    return chunks

def _extract_embedding_values(resp) -> list:
    emb = getattr(resp, "embedding", None)
    if emb is not None:
        values = getattr(emb, "values", emb)
        if isinstance(values, (list, tuple)):
            return list(values)
        if isinstance(values, dict) and "values" in values:
            return list(values["values"])
    if isinstance(resp, dict):
        emb = resp.get("embedding")
        if isinstance(emb, dict) and "values" in emb:
            return list(emb["values"])
        if isinstance(emb, (list, tuple)):
            return list(emb)
    return []

def _embed_text(text: str) -> np.ndarray:
    try:
        resp = genai.embed_content(model=EMBEDDING_MODEL, content=text)
        values = _extract_embedding_values(resp)
        if not values:
            return np.array([])
        return np.array(values, dtype=np.float32)
    except Exception:
        return np.array([])

def _build_retrieval_index(full_text: str):
    text_hash = str(hash(full_text))
    if (st.session_state.get("cached_text_hash") == text_hash and 
        st.session_state.get("retrieval_embeddings") is not None):
        return
    chunks = _chunk_text(full_text)
    embeddings = []

    max_chunks = min(len(chunks), 10)
    chunks = chunks[:max_chunks]
    
    for ch in chunks:
        try:
            vec = _embed_text(ch)
            if vec.size == 0:
                embeddings = []
                break
            embeddings.append(vec)
        except Exception as e:
            if "429" in str(e) or "quota" in str(e).lower():
                st.warning("⚠️ Gemini embedding quota exceeded. Using full-context fallback.")
                embeddings = []
                break
            else:
                embeddings = []
                break
    
    if embeddings:
        matrix = np.vstack(embeddings)
        st.session_state.retrieval_chunks = chunks
        st.session_state.retrieval_embeddings = matrix
        st.session_state.retrieval_norms = np.linalg.norm(matrix, axis=1, keepdims=True) + 1e-10
        st.session_state.cached_text_hash = text_hash
    else:
        st.session_state.retrieval_chunks = None
        st.session_state.retrieval_embeddings = None
        st.session_state.retrieval_norms = None
        st.session_state.cached_text_hash = None

def _retrieve_top_k(query: str, k: int = 5) -> list:
    if not query:
        return []
    chunks = st.session_state.get("retrieval_chunks")
    emb = st.session_state.get("retrieval_embeddings")
    norms = st.session_state.get("retrieval_norms")
    if not chunks or emb is None or norms is None:
        return []
    q_vec = _embed_text(query)
    if q_vec.size == 0:
        return []
    q_vec = q_vec.reshape(1, -1)
    q_norm = np.linalg.norm(q_vec, axis=1, keepdims=True) + 1e-10
    sims = (emb @ q_vec.T) / (norms * q_norm)
    sims = sims.ravel()
    top_idx = np.argsort(-sims)[: max(1, k)]
    return [chunks[i] for i in top_idx]

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        text = "\n".join(parts).strip()
        return text
    except Exception:
        return ""

# --------------------- DOC/DOCX extractors -------------------------
def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(docx_bytes)
            tmp_path = tmp.name
        doc = Document(tmp_path)
        parts = []

        # paragraphs
        parts.extend(p.text for p in doc.paragraphs if p.text)

        # tables (baris → Markdown)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                parts.append("| " + " | ".join(cells) + " |")

        return "\n".join(parts).strip()
    except Exception:
        return ""
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

def extract_text_from_doc_bytes(doc_bytes: bytes) -> str:
    try:
        with tempfile.TemporaryDirectory() as td:
            src = os.path.join(td, "in.doc")
            with open(src, "wb") as f:
                f.write(doc_bytes)

            # unoconv -> txt
            try:
                subprocess.run(["unoconv", "-f", "txt", src], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                txt_path = src.replace(".doc", ".txt")
                if os.path.exists(txt_path):
                    with open(txt_path, "r", encoding="utf-8", errors="ignore") as t:
                        return t.read()
            except Exception:
                pass

            # soffice -> txt
            try:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "txt:Text", "--outdir", td, src],
                    check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
                )
                base = os.path.splitext(os.path.basename(src))[0]
                txt_path = os.path.join(td, base + ".txt")
                if os.path.exists(txt_path):
                    with open(txt_path, "r", encoding="utf-8", errors="ignore") as t:
                        return t.read()
            except Exception:
                pass

            # soffice -> pdf -> extractor PDF / OCR Gemini
            try:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "pdf", "--outdir", td, src],
                    check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE
                )
                pdf_path = os.path.join(td, "in.pdf")
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as pf:
                        pdf_bytes = pf.read()
                    txt = extract_text_from_pdf_bytes(pdf_bytes)
                    if txt and len(txt) > 30:
                        return txt
                    return gemini_ocr_pdf(pdf_bytes, filename="converted_from_doc.pdf")
            except Exception:
                pass

            # strings (last resort)
            try:
                out = subprocess.run(["strings", src], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                txt = out.stdout.decode("utf-8", errors="ignore")
                return txt
            except Exception:
                return ""
    except Exception:
        return ""

# --------------------- OCR helpers -------------------------
def gemini_ocr_image(image_bytes: bytes) -> str:
    img = Image.open(io.BytesIO(image_bytes))
    prompt_doc = (
        "Convert this document image into clean Markdown. "
        "Preserve headings, lists, and tables (use Markdown tables). "
        "Maintain natural reading order."
    )
    text = _generate_with_fallback([prompt_doc, img])
    if not text or len(text.strip()) < 30:
        prompt_cap_id = (
            "Jelaskan gambar ini secara ringkas, jelas, dan akurat. "
            "Sebutkan objek utama, konteks, warna, teks (jika ada), dan hal penting lainnya."
        )
        prompt_cap_en = (
            "Describe this image concisely and accurately. "
            "Mention main objects, context, colors, any visible text, and other important details."
        )
        prompt_cap = prompt_cap_id if ANSWER_LANGUAGE == "Bahasa Indonesia" else prompt_cap_en
        text = _generate_with_fallback([prompt_cap, img])
    return text

def gemini_ocr_pdf(pdf_bytes: bytes, filename: str = "upload.pdf") -> str:
    suffix = os.path.splitext(filename)[1] or ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name
    try:
        file_obj = genai.upload_file(
            path=tmp_path,
            mime_type="application/pdf",
            display_name=filename
        )
        try:
            for _ in range(30):
                f = genai.get_file(file_obj.name)
                state = getattr(getattr(f, "state", None), "name", getattr(f, "state", ""))
                if str(state).upper() == "ACTIVE":
                    break
                time.sleep(1)
            else:
                return "Error: File processing timed out. Please try again."
        except Exception:
            time.sleep(2)
        prompt = (
            "Extract the full content of this PDF as clean Markdown. "
            "Preserve headings and tables. If pages are scanned, perform OCR first."
        )
        return _generate_with_fallback([file_obj, prompt])
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

# --------------------- PDF/Image/DOC/DOCX pipeline -------------------------
def process_document_with_gemini(kind: str, name: str, data: bytes) -> str:
    if kind == "pdf":
        # 1) coba PyPDF2
        raw_text = extract_text_from_pdf_bytes(data)
        # deteksi sederhana apakah terlihat seperti tabel markdown
        looks_like_table = bool(re.search(r'\|\s*[^|]+\s*\|', raw_text))
        if len(raw_text) >= 200 and looks_like_table:
            return raw_text
        # 2) jika tidak terlihat tabel (atau pendek), paksa OCR Gemini (lebih bagus untuk tabel)
        return gemini_ocr_pdf(data, filename=name)

    if kind == "image":
        return gemini_ocr_image(data)

    if kind == "docx":
        text = extract_text_from_docx_bytes(data)
        if text and len(text) >= 50:
            return text
        # fallback minimal (tidak sebaik parser native)
        return _generate_with_fallback([
            "Extract the full content of this Office document as clean Markdown.",
            data
        ])

    if kind == "doc":
        text = extract_text_from_doc_bytes(data)
        if text and len(text) >= 50:
            return text
        return "No content extracted from .doc (older Word). Please ensure LibreOffice/unoconv is installed for better results."

    return "Unsupported document kind."

def answer_from_image(image_bytes: bytes, question: str) -> str:
    try:
        img = Image.open(io.BytesIO(image_bytes))
        if ANSWER_LANGUAGE == "Bahasa Indonesia":
            prompt = (
                "Anda adalah asisten analisis visual. Jawab pertanyaan pengguna hanya berdasarkan gambar ini.\n"
                "Jika informasi tidak terlihat pada gambar, katakan tidak ada.\n"
                f"Pertanyaan: {question}"
            )
        else:
            prompt = (
                "You are a visual analysis assistant. Answer the user's question based only on this image.\n"
                "If the information is not visible in the image, say so.\n"
                f"Question: {question}"
            )
        return _generate_with_fallback([prompt, img]) or "No response text."
    except Exception as e:
        return f"Error generating visual answer: {e}"

def generate_response(context: str, query: str) -> str:
    if not context or len(context) < 10:
        if "image_bytes" in st.session_state and isinstance(st.session_state.image_bytes, dict):
            for img_name, img_bytes in st.session_state.image_bytes.items():
                if any(word.lower() in img_name.lower() for word in query.lower().split()):
                    return answer_from_image(img_bytes, query)
            first_img = next(iter(st.session_state.image_bytes.values()))
            return answer_from_image(first_img, query)
    try:
        retrieved_chunks = _retrieve_top_k(query, k=5)
        if retrieved_chunks:
            context_block = "\n\n---\n\n".join(retrieved_chunks)
        else:
            context_block = context
        
        doc_context = ""
        if st.session_state.get("documents") and len(st.session_state.documents) > 1:
            doc_names = [doc['name'] for doc in st.session_state.documents]
            doc_context = f"\n\nAvailable documents: {', '.join(doc_names)}"
        
        if ANSWER_LANGUAGE == "Bahasa Indonesia":
            prompt = f"""
Anda adalah asisten analisis dokumen. Gunakan konteks berikut untuk menjawab:

{context_block}{doc_context}

Pertanyaan pengguna:
{query}

Jawab dalam Bahasa Indonesia secara ringkas, jelas, dan akurat. Jika jawabannya tidak terdapat pada konteks, katakan: "Tidak ditemukan pada dokumen". Jika ada beberapa dokumen, sebutkan dari dokumen mana informasi berasal.
"""
        else:
            prompt = f"""
You are a document analysis assistant. Use only the context below to answer:

{context_block}{doc_context}

User question:
{query}

Respond in English concisely. If the answer is not in the context, say so. If there are multiple documents, mention which document the information comes from.
"""
        return _generate_with_fallback(prompt) or "No response text."
    except Exception as e:
        return f"Error generating response: {e}"

# --------------------- API Fallback & Caching ---------------------
def generate_with_mistral_fallback(prompt: str) -> str:
    if not mistral_client:
        return "Error: No Mistral API key configured for fallback."
    try:
        response = mistral_client.chat(
            model="mistral-large-latest",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error with Mistral fallback: {e}"

def generate_response_with_fallback(context: str, query: str) -> str:
    try:
        return generate_response(context, query)
    except Exception as e:
        if "429" in str(e) or "quota" in str(e).lower():
            st.warning("⚠️ Gemini API quota exceeded. Falling back to Mistral API...")
            if ANSWER_LANGUAGE == "Bahasa Indonesia":
                mistral_prompt = f"""
Anda adalah asisten analisis dokumen. Gunakan konteks berikut untuk menjawab:

{context}

Pertanyaan pengguna:
{query}

Jawab dalam Bahasa Indonesia secara ringkas, jelas, dan akurat. Jika jawabannya tidak terdapat pada konteks, katakan: "Tidak ditemukan pada dokumen".
"""
            else:
                mistral_prompt = f"""
You are a document analysis assistant. Use only the context below to answer:

{context}

User question:
{query}

Respond in English concisely. If the answer is not in the context, say so.
"""
            return generate_with_mistral_fallback(mistral_prompt)
        else:
            return f"Error generating response: {e}"

# --------------------- Document Management Helpers ---------------------
def clear_all_document_state():
    st.session_state.documents = []
    st.session_state.ocr_content = None
    st.session_state.retrieval_chunks = None
    st.session_state.retrieval_embeddings = None
    st.session_state.retrieval_norms = None
    st.session_state.image_bytes = {}
    st.session_state.chat_history = []

def rebuild_document_content():
    if st.session_state.documents:
        all_content = [f"--- DOCUMENT: {d['name']} ---\n{d['content']}" for d in st.session_state.documents]
        st.session_state.ocr_content = "\n\n".join(all_content)
        _build_retrieval_index(st.session_state.ocr_content)
    else:
        st.session_state.ocr_content = None
        st.session_state.retrieval_chunks = None
        st.session_state.retrieval_embeddings = None
        st.session_state.retrieval_norms = None


# ======================= DASHBOARD HELPERS =======================
_ID_STOPWORDS = {
    "yang","dan","di","ke","dari","untuk","pada","dengan","ini","itu","ada","tidak","atau","karena","sebagai",
    "dalam","atas","oleh","sebuah","para","akan","juga","sudah","belum","saat","kami","kita","mereka","anda",
    "ia","dia","tersebut","rp","usd","pt","tbk","persero","co","ltd","inc"
}
_EN_STOPWORDS = {
    "the","and","of","to","in","for","on","at","by","with","from","as","is","are","was","were","be","been","a","an",
    "this","that","these","those","it","its","we","you","they","he","she","them","our","your","their",
    "or","not","but","if","then","so","than","such","per","vs"
}
STOPWORDS = _ID_STOPWORDS | _EN_STOPWORDS

def _tokenize(text: str) -> list:
    text = text.lower()
    text = text.translate(str.maketrans({c: " " for c in string.punctuation}))
    toks = [t for t in text.split() if t and t not in STOPWORDS and not t.isdigit() and len(t) > 2]
    return toks

def _compute_ocr_quality(text: str) -> float:
    if not text: return 0.0
    n = len(text)
    good = sum(ch.isalnum() or ch.isspace() or ch in ".,:;-%()[]|/+\n" for ch in text)
    bad = text.count(" ")
    short_lines = sum(1 for ln in text.splitlines() if 0 < len(ln.strip()) < 3)
    score = (good / n) * 100.0
    score -= min(25, bad * 0.5)
    score -= min(15, short_lines * 0.2)
    return max(0.0, min(100.0, score))

def _structure_stats(md_text: str) -> dict:
    if not md_text:
        return {"text": 0, "tables": 0, "images": 0}
    lines = md_text.splitlines()
    table_lines = sum(1 for ln in lines if ln.strip().startswith("|") and ln.count("|") >= 2)
    image_tags = md_text.count("![") + md_text.lower().count("<img")
    text_chars = len(md_text)
    return {"text": text_chars, "tables": table_lines, "images": image_tags}

def _extract_sections(md_text: str):
    sections = []
    current_title = "Intro"
    current_buf = []
    for ln in md_text.splitlines():
        if ln.startswith("--- DOCUMENT:"):
            if current_buf:
                sections.append((current_title, "\n".join(current_buf).strip()))
                current_buf = []
            current_title = ln.replace("--- DOCUMENT:", "").strip()
        elif re.match(r"^\s{0,3}#{1,3}\s+\S", ln):
            if current_buf:
                sections.append((current_title, "\n".join(current_buf).strip()))
                current_buf = []
            current_title = re.sub(r"^\s{0,3}#{1,3}\s+", "", ln).strip()
        else:
            current_buf.append(ln)
    if current_buf:
        sections.append((current_title, "\n".join(current_buf).strip()))
    return sections[:12] if sections else [("All", md_text)]

def _missing_matrix(sections):
    cats = ["N/A/NA", "Empty Lines", "Dashes(-/—)", "Question(?)"]
    labels = [title[:40] + ("…" if len(title) > 40 else "") for title,_ in sections]
    matrix = []
    for _, txt in sections:
        lines = txt.splitlines()
        na = sum(bool(re.search(r"\b(n/?a|tidak tersedia|kosong)\b", ln, re.I)) for ln in lines)
        empty = sum(1 for ln in lines if not ln.strip())
        dashes = sum(ln.count("-") + ln.count("—") for ln in lines)
        ques = txt.count("?")
        matrix.append([na, empty, dashes, ques])
    return labels, cats, np.array(matrix).T

def _is_financial_report(text: str) -> bool:
    keys = ["revenue","pendapatan","penjualan","income","profit","laba","rugi",
            "neraca","balance sheet","arus kas","cash flow","laba kotor","gross","operating","net"]
    t = text.lower()
    return any(k in t for k in keys)

_num_pat = re.compile(r"([\-]?\d{1,3}(?:[\.,]\d{3})*(?:[\.,]\d{1,2})?)")

def _parse_number(s: str):
    s = s.strip()
    if not s: return None
    if "." in s and "," in s:
        if s.find(".") < s.find(","):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    else:
        if "," in s and "." not in s:
            s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        digits = re.sub(r"[^\d\-\.]", "", s)
        try:
            return float(digits)
        except Exception:
            return None

def _extract_financials(text: str):
    rev = {}
    prof = {}
    gross = {}
    op = {}
    net = {}
    assets = equity = curr_assets = curr_liab = debt = None

    lines = text.splitlines()
    for ln in lines:
        years = re.findall(r"\b(20\d{2}|19\d{2})\b", ln)
        if not years:
            continue
        y = int(years[0])

        if re.search(r"\b(revenue|pendapatan|penjualan)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    rev[y] = val
        if re.search(r"\b(net income|laba bersih|laba/rugi bersih|profit|laba)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    prof[y] = val
        if re.search(r"\b(gross profit|laba kotor)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    gross[y] = val
        if re.search(r"\b(operating profit|laba usaha|laba operasi)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    op[y] = val
        if re.search(r"\b(net profit|net income|laba bersih)\b", ln, re.I):
            m = _num_pat.search(ln)
            if m:
                val = _parse_number(m.group(1))
                if val is not None:
                    net[y] = val

        if assets is None and re.search(r"\b(total assets|jumlah aset)\b", ln, re.I):
            m = _num_pat.search(ln); assets = _parse_number(m.group(1)) if m else None
        if equity is None and re.search(r"\b(total equity|ekuitas)\b", ln, re.I):
            m = _num_pat.search(ln); equity = _parse_number(m.group(1)) if m else None
        if curr_assets is None and re.search(r"\b(current assets|aset lancar)\b", ln, re.I):
            m = _num_pat.search(ln); curr_assets = _parse_number(m.group(1)) if m else None
        if curr_liab is None and re.search(r"\b(current liab|liabilitas lancar|utang lancar)\b", ln, re.I):
            m = _num_pat.search(ln); curr_liab = _parse_number(m.group(1)) if m else None
        if debt is None and re.search(r"\b(total debt|utang|pinjaman)\b", ln, re.I):
            m = _num_pat.search(ln); debt = _parse_number(m.group(1)) if m else None

    return {
        "revenue_by_year": dict(sorted(rev.items())),
        "profit_by_year": dict(sorted(prof.items())),
        "gross_by_year": dict(sorted(gross.items())),
        "operating_by_year": dict(sorted(op.items())),
        "net_by_year": dict(sorted(net.items())),
        "assets": assets, "equity": equity,
        "current_assets": curr_assets, "current_liabilities": curr_liab, "debt": debt
    }

def _yoy_growth(series: dict) -> dict:
    ys = sorted(series.keys())
    growth = {}
    for i in range(1, len(ys)):
        y0, y1 = ys[i-1], ys[i]
        if series[y0] and series[y0] != 0:
            growth[y1] = (series[y1] - series[y0]) / abs(series[y0]) * 100.0
    return growth

def _readability_from_avg_sentence_len(text: str):
    if not text: return 0.0, 0.0
    sents = re.split(r"[.!?\n]+", text)
    sents = [s.strip() for s in sents if s.strip()]
    if not sents: return 0.0, 0.0
    words = sum(len(s.split()) for s in sents)
    avg = words / len(sents)
    score = 100.0 - ((avg - 8) / (40 - 8)) * 100.0
    score = max(0.0, min(100.0, score))
    return avg, score

def _ner_counts_with_gemini(text: str):
    if not google_api_key:
        return None
    try:
        prompt = """
Extract counts of named entities by coarse type from the text.
Only return a compact JSON like {"ORG": 12, "PERSON": 5, "LOC": 7}. Types limited to ORG, PERSON, LOC.
If none detected, use 0.
Text:
""" + _truncate_context(text, 20000)
        raw = _generate_with_fallback(prompt)
        data = json.loads(re.findall(r"\{.*\}", raw, re.S)[0])
        return {"ORG": int(data.get("ORG", 0)), "PERSON": int(data.get("PERSON", 0)), "LOC": int(data.get("LOC", 0))}
    except Exception:
        return None

def _ner_counts_naive(text: str):
    ORG = len(re.findall(r"\b(PT|Tbk|Persero|Inc\.?|Ltd\.?|LLC|Corp\.?)\b", text))
    PERSON = len(re.findall(r"\b[A-Z][a-z]+ [A-Z][a-z]+(?: [A-Z][a-z]+)?\b", text))
    LOC = len(re.findall(r"\b(Jakarta|Bandung|Surabaya|Medan|Indonesia|Singapore|Malaysia|USA|Europe|Asia)\b", text))
    return {"ORG": ORG, "PERSON": PERSON, "LOC": LOC}


# ======================= TABLE PARSER & AVG INTENT =======================
_TABLE_ROW_RE = re.compile(r'^\s*\|(.+)\|\s*$')
_TABLE_SEP_RE = re.compile(r'^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$')

def _parse_markdown_tables(md_text: str) -> list:
    """
    Mengembalikan list of (DataFrame, meta) dari tabel Markdown di teks.
    Meta: {"doc_section": str, "table_index": int}
    Mencoba infer header saat tidak ada header eksplisit.
    """
    tables = []
    lines = md_text.splitlines()
    buf = []

    def _flush(buf_lines, t_idx):
        if not buf_lines:
            return
        rows = []
        for ln in buf_lines:
            m = _TABLE_ROW_RE.match(ln)
            if not m:
                continue
            cells = [c.strip() for c in m.group(1).split('|')]
            rows.append(cells)

        if not rows:
            return

        # Header detection:
        header = None
        if len(rows) >= 2 and _TABLE_SEP_RE.match(buf_lines[1]):  # header | sep | data...
            header = [h if h else f"col_{i+1}" for i, h in enumerate(rows[0])]
            data = rows[2:] if len(rows) > 2 else []
        else:
            # Coba infer: jika baris pertama lebih "tekstual" → header
            first = rows[0]
            nonnum = sum(1 for x in first if not re.fullmatch(r'[-+]?\d+([.,]\d+)?', (x or '')))
            if nonnum >= max(1, len(first) // 2):
                header = [h if h else f"col_{i+1}" for i, h in enumerate(first)]
                data = rows[1:]
            else:
                ncol = len(first)
                header = [f"col_{i+1}" for i in range(ncol)]
                data = rows

        # Normalisasi panjang baris
        maxc = len(header)
        norm_rows = []
        for r in data:
            if len(r) < maxc:
                r = r + [''] * (maxc - len(r))
            elif len(r) > maxc:
                r = r[:maxc]
            norm_rows.append(r)

        try:
            df = pd.DataFrame(norm_rows, columns=header)
            tables.append((df, {"doc_section": "All", "table_index": t_idx}))
        except Exception:
            pass

    t_idx = 0
    for ln in lines:
        if _TABLE_ROW_RE.match(ln):
            buf.append(ln)
        else:
            if buf:
                _flush(buf, t_idx)
                t_idx += 1
                buf = []
    if buf:
        _flush(buf, t_idx)

    return tables

def _is_avg_intent(q: str) -> bool:
    ql = q.lower()
    keys = ["rata rata", "rata-rata", "average", "avg", "mean"]
    return any(k in ql for k in keys)

def _extract_target_column_terms(q: str) -> list:
    # Dalam tanda kutip
    quoted = re.findall(r'["“”](.+?)["“”]', q)
    terms = [t.strip() for t in quoted if t.strip()]
    # Setelah kata kunci
    m = re.search(r'\b(kolom|column|field)\s+([A-Za-z0-9_\-\s]+)', q, re.I)
    if m:
        chunk = m.group(2).strip()
        chunk = re.split(r'[,.;:?]|rata|average|mean', chunk, 1)[0].strip()
        if chunk:
            terms.append(chunk)
    # Fallback
    toks = _tokenize(q)
    toks = [t for t in toks if t not in {"rata", "rata-rata", "average", "mean", "avg", "nilai", "value"}]
    if toks:
        terms.append(max(toks, key=len))
    seen = set(); res = []
    for t in terms:
        tt = t.lower()
        if tt not in seen:
            seen.add(tt); res.append(t)
    return res[:3]

def _normalize_col(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r'\s+', ' ', s)
    s = s.replace('_', ' ')
    return s

def _guess_best_column(df: pd.DataFrame, query: str) -> str | None:
    if df.empty or df.shape[1] == 0:
        return None
    colmap = {c: _normalize_col(c) for c in df.columns}
    terms = _extract_target_column_terms(query)
    synonyms = {
        "harga": ["price", "amount", "nilai", "nominal", "jumlah"],
        "pendapatan": ["revenue", "sales", "omzet"],
        "laba": ["profit", "net profit", "laba bersih", "income"],
        "tanggal": ["date", "periode", "period", "bulan", "month", "year", "tahun"],
        "qty": ["quantity", "jumlah", "kuantitas", "volume"],
    }
    candidates = []
    for t in terms:
        nt = _normalize_col(t)
        for orig, norm in colmap.items():
            if nt == norm or nt in norm or norm in nt:
                candidates.append(orig)
        close = difflib.get_close_matches(nt, list(colmap.values()), n=1, cutoff=0.75)
        if close:
            for orig, norm in colmap.items():
                if norm == close[0]:
                    candidates.append(orig)
        for base, syns in synonyms.items():
            if nt == base or nt in syns:
                close2 = difflib.get_close_matches(base, list(colmap.values()), n=1, cutoff=0.6)
                if close2:
                    for orig, norm in colmap.items():
                        if norm == close2[0]:
                            candidates.append(orig)
    for c in candidates:
        s = pd.to_numeric(
            df[c].astype(str)
                .str.replace(r'[^\d\-\.,]', '', regex=True)
                .str.replace('.', '', regex=False)
                .str.replace(',', '.', regex=False),
            errors='coerce'
        )
        if s.notna().sum() >= max(2, int(len(s)*0.5)):
            return c
    best_col = None; best_score = -1
    for c in df.columns:
        s = pd.to_numeric(
            df[c].astype(str)
                .str.replace(r'[^\d\-\.,]', '', regex=True)
                .str.replace('.', '', regex=False)
                .str.replace(',', '.', regex=False),
            errors='coerce'
        )
        score = s.notna().sum()
        if score > best_score:
            best_score = score; best_col = c
    return best_col

# --------- Fallback parser nilai dari teks polos ---------
def _to_float_id_en(num: str) -> float | None:
    if not num:
        return None
    s = num.strip()
    if "." in s and "," in s:
        if s.find(".") < s.find(","):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    else:
        if "," in s and "." not in s:
            s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        try:
            digits = re.sub(r"[^\d\.\-]", "", s)
            return float(digits) if digits else None
        except Exception:
            return None

def _extract_scores_from_text(text: str) -> list[float]:
    """
    Fallback untuk dokumen ranking seperti contoh:
    '... 22090098 NAMA 93,7 089670916052 ...'
    Strategi:
      1) Cari angka 0..100 yang diikuti nomor HP -> (score, phone)
      2) Atau pola 'rank nim nama score' -> ambil score
      3) Atau angka di dekat kata 'nilai'
    """
    scores: list[float] = []
    t = " ".join(text.split())  # rapikan spasi

    # 1) Angka sebelum nomor HP (0xxxxxxxxxx…)
    p_phone = re.compile(r'(\d{1,3}(?:[.,]\d{1,2})?)\s+(0\d{9,13})')
    for m in p_phone.finditer(t):
        val = _to_float_id_en(m.group(1))
        if val is not None and 0 <= val <= 100:
            scores.append(val)

    # 2) 'rank nim nama score'
    p_rank = re.compile(
        r'\b\d{1,3}\s+(?:\d{8,12})\s+[A-Za-zÀ-ÿ\.\'\- ]{2,}?'
        r'\s+(\d{1,3}(?:[.,]\d{1,2})?)\b'
    )
    for m in p_rank.finditer(t):
        val = _to_float_id_en(m.group(1))
        if val is not None and 0 <= val <= 100:
            scores.append(val)

    # 3) Di dekat kata "nilai"
    if not scores:
        near = []
        for m in re.finditer(r'(nilai[^0-9]{0,30})(\d{1,3}(?:[.,]\d{1,2})?)', t, flags=re.I):
            val = _to_float_id_en(m.group(2))
            if val is not None and 0 <= val <= 100:
                near.append(val)
        if len(near) >= 3:
            scores.extend(near)

    # Dedup ringan
    if scores:
        seen = set(); uniq = []
        for v in scores:
            key = round(v, 2)
            if key not in seen:
                seen.add(key); uniq.append(v)
        scores = uniq

    return scores

# ========== BARU: pilih dokumen berdasar pertanyaan ==========
def _normalize_filename(name: str) -> str:
    base = os.path.splitext(name)[0]
    base = re.sub(r'[^a-zA-Z0-9]+', ' ', base).strip().lower()
    return base

def select_docs_for_query(query: str, documents: list) -> list:
    """
    Jika user menyebutkan nama file (tanpa ekstensi) pada pertanyaan,
    kembalikan hanya dokumen yang disebut. Jika tidak ada yang match, kembalikan semua.
    """
    q = query.lower()
    matches = []
    for d in documents:
        norm = _normalize_filename(d["name"])
        # match frasa penuh atau sebagian (split ke kata-kata unik)
        if norm and norm in q:
            matches.append(d)
        else:
            tokens = [t for t in norm.split() if len(t) > 2]
            hit = sum(1 for t in tokens if t in q)
            if hit >= max(1, len(tokens)//2):
                matches.append(d)
    if len(matches) == 1:
        return matches
    if len(matches) >= 2:
        return matches
    return documents

# ========== MODIFIED: hitung rata-rata per dokumen ==========
def compute_average_per_doc(query: str, documents: list) -> list[dict]:
    """
    Hitung rata-rata kolom target di tiap dokumen.
    Return list of dict: [{'doc':..., 'column':..., 'value':..., 'n':...}, ...]
    """
    if not _is_avg_intent(query):
        return []

    results = []

    for doc in documents:
        content = doc.get("content") or ""
        tables = _parse_markdown_tables(content)

        best = None

        # --- cari di tabel markdown ---
        for df, meta in tables:
            if df.empty:
                continue
            col = _guess_best_column(df, query)
            if not col:
                continue
            s = pd.to_numeric(
                df[col].astype(str)
                    .str.replace(r'[^\d\-\.,]', '', regex=True)
                    .str.replace('.', '', regex=False)
                    .str.replace(',', '.', regex=False),
                errors='coerce'
            )
            vals = s.dropna()
            if len(vals) == 0:
                continue
            mean_val = float(vals.mean())
            cand = {"doc": doc["name"], "column": col,
                    "value": mean_val, "n": int(len(vals))}
            if (best is None) or (cand["n"] > best["n"]):
                best = cand

        # --- fallback dari teks polos ---
        if best is None or best["n"] < 3:
            scores = _extract_scores_from_text(content)
            if len(scores) >= 3:
                mean_val = float(np.mean(scores))
                cand = {"doc": doc["name"], "column": "Nilai (fallback)",
                        "value": mean_val, "n": int(len(scores))}
                best = cand

        if best:
            results.append(best)

    return results


# --------------------- UI Layout -----------------------
col1, col2 = st.columns([1, 1])

with col1:
    st.header("Document Upload")
    uploaded_files = st.file_uploader(
        "Upload multiple documents (DOC, DOCX, PDF, PNG, JPG, JPEG)",
        type=["pdf", "png", "jpg", "jpeg", "doc", "docx"],
        accept_multiple_files=True,
    )
    url_input = st.text_input("Or enter a URL (web page or document):")
    st.session_state["url_input"] = url_input

    process_button = st.button("Process Documents")

    if "documents" not in st.session_state:
        st.session_state.documents = []
    if "ocr_content" not in st.session_state:
        st.session_state.ocr_content = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    # --------------------- URL processing helper -----------------------
    def process_url_to_content(url: str) -> tuple:
        try:
            r = requests.get(
                url,
                timeout=30,
                headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                },
                allow_redirects=True,
            )
            r.raise_for_status()
            content_type = r.headers.get("Content-Type", "").lower()
            clean_url = url.split("?")[0]
            ext = os.path.splitext(clean_url)[1].lower()
            data = r.content

            # signature
            is_pdf_sig = data[:4] == b"%PDF"
            is_png_sig = data[:8] == b"\x89PNG\r\n\x1a\n"
            is_jpg_sig = data[:3] == b"\xff\xd8\xff"

            # Detect Office CT
            is_docx_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in content_type
            is_doc_ct = "application/msword" in content_type

            chosen_kind = None
            if is_pdf_sig or "pdf" in content_type or ext == ".pdf":
                chosen_kind = "pdf"
            elif is_png_sig or is_jpg_sig or any(img_ct in content_type for img_ct in ["image/png", "image/jpeg", "image/jpg"]) or ext in [".png", ".jpg", ".jpeg"]:
                chosen_kind = "image"
            elif ext == ".docx" or is_docx_ct:
                chosen_kind = "docx"
            elif ext == ".doc" or is_doc_ct:
                chosen_kind = "doc"

            if not chosen_kind and content_type.startswith("text/html"):
                html_text = r.text
                links = re.findall(
                    r'href=["\']([^"\']+\.(?:pdf|png|jpe?g|docx?|DOCX?))(?:[#\?][^"\']*)?["\']',
                    html_text, flags=re.IGNORECASE
                )
                if links:
                    target_url = urljoin(url, links[0])
                    rr = requests.get(
                        target_url,
                        timeout=30,
                        headers={
                            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                        },
                        allow_redirects=True,
                    )
                    rr.raise_for_status()
                    target_ct = rr.headers.get("Content-Type", "").lower()
                    tdata = rr.content
                    t_clean = target_url.split("?")[0]
                    t_ext = os.path.splitext(t_clean)[1].lower()

                    if tdata[:4] == b"%PDF" or "pdf" in target_ct or t_ext == ".pdf":
                        chosen_kind = "pdf"; data = tdata; clean_url = t_clean
                    elif tdata[:8] == b"\x89PNG\r\n\x1a\n" or tdata[:3] == b"\xff\xd8\xff" or any(ic in target_ct for ic in ["image/png","image/jpeg","image/jpg"]) or t_ext in [".png",".jpg",".jpeg"]:
                        chosen_kind = "image"; data = tdata; clean_url = t_clean
                        st.session_state.image_bytes = data
                    elif t_ext == ".docx" or "vnd.openxmlformats-officedocument.wordprocessingml.document" in target_ct:
                        chosen_kind = "docx"; data = tdata; clean_url = t_clean
                    elif t_ext == ".doc" or "application/msword" in target_ct:
                        chosen_kind = "doc"; data = tdata; clean_url = t_clean

                if not chosen_kind:
                    stripped = re.sub(r"<script[\s\S]*?</script>", " ", html_text, flags=re.IGNORECASE)
                    stripped = re.sub(r"<style[\s\S]*?</style>", " ", stripped, flags=re.IGNORECASE)
                    text_only = re.sub(r"<[^>]+>", " ", stripped)
                    text_only = re.sub(r"\s+", " ", text_only).strip()
                    st.session_state.ocr_content = text_only
                    if st.session_state.ocr_content:
                        _build_retrieval_index(st.session_state.ocr_content)
                        return True, "Webpage processed as text."
                    return False, "No content extracted from webpage."

            if not chosen_kind:
                return False, f"Unsupported content type: {content_type or ext}"

            st.session_state.ocr_content = process_document_with_gemini(
                chosen_kind, os.path.basename(clean_url) or "download", data
            )
            if chosen_kind == "image":
                st.session_state.image_bytes = data
            if st.session_state.ocr_content:
                _build_retrieval_index(st.session_state.ocr_content)
                return True, "Document processed successfully!"
            return False, "No content extracted."
        except Exception as e:
            return False, f"Error processing document: {e}"

    if process_button:
        if not google_api_key:
            st.error("Please provide a valid Google API Key for OCR/processing.")
        if uploaded_files:
            with st.spinner("Processing documents..."):
                all_content = []
                for uploaded_file in uploaded_files:
                    try:
                        ext = os.path.splitext(uploaded_file.name)[1].lower()
                        if ext == ".pdf":
                            kind = "pdf"
                        elif ext in [".png", ".jpg", ".jpeg"]:
                            kind = "image"
                        elif ext == ".docx":
                            kind = "docx"
                        elif ext == ".doc":
                            kind = "doc"
                        else:
                            kind = "unknown"

                        if kind == "image":
                            if "image_bytes" not in st.session_state:
                                st.session_state.image_bytes = {}
                            st.session_state.image_bytes[uploaded_file.name] = uploaded_file.getvalue()
                        
                        if kind == "unknown":
                            st.warning(f"Unsupported file type for {uploaded_file.name}. Skipped.")
                            continue

                        content = process_document_with_gemini(
                            kind, uploaded_file.name, uploaded_file.getvalue()
                        )
                        
                        if content:
                            doc_info = {
                                "name": uploaded_file.name,
                                "type": kind,
                                "content": content,
                                "size": len(uploaded_file.getvalue())
                            }
                            st.session_state.documents.append(doc_info)
                            all_content.append(f"--- DOCUMENT: {uploaded_file.name} ---\n{content}")
                            st.success(f"Document {uploaded_file.name} processed successfully!")
                        else:
                            st.warning(f"No content extracted from {uploaded_file.name}.")
                    except Exception as e:
                        st.error(f"Error processing {uploaded_file.name}: {e}")
                
                if all_content:
                    st.session_state.ocr_content = "\n\n".join(all_content)
                    _build_retrieval_index(st.session_state.ocr_content)
                    st.success(f"All {len(uploaded_files)} documents processed and combined!")
        if url_input:
            with st.spinner("Downloading & processing from URL..."):
                success, msg = process_url_to_content(url_input)
                if success:
                    if st.session_state.get("ocr_content"):
                        url_doc_info = {
                            "name": f"URL: {url_input[:50]}...",
                            "type": "url",
                            "content": st.session_state.ocr_content,
                            "size": len(st.session_state.ocr_content)
                        }
                        st.session_state.documents.append(url_doc_info)
                    st.success(msg)
                else:
                    st.error(msg)
        if not uploaded_files and not url_input:
            st.warning("Please upload a document or provide a URL.")

with col2:
    st.header("Document Q&A")

    if st.session_state.documents:
        st.markdown(f"**{len(st.session_state.documents)} document(s) loaded.**")
        
        total_chars = sum(doc['size'] for doc in st.session_state.documents)
        total_mb = total_chars / (1024 * 1024)
        
        c1, c2, c3 = st.columns([2, 2, 1])
        with c1:
            st.metric("Total Documents", len(st.session_state.documents))
        with c2:
            st.metric("Total Content", f"{total_chars:,} chars")
        with c3:
            st.metric("Memory Usage", f"{total_mb:.1f} MB")
        
        if len(st.session_state.documents) > 10:
            st.warning("⚠️ Many documents loaded. Consider removing some to improve performance.")
        elif total_chars > 500000:
            st.warning("⚠️ Very large document collection. Consider removing some documents to avoid processing limits.")
        elif total_chars > 200000 and len(st.session_state.documents) > 3:
            st.warning("⚠️ Large document collection. Consider removing some documents to avoid processing limits.")
        
        with st.expander("📋 Document List"):
            for i, doc in enumerate(st.session_state.documents):
                cc1, cc2 = st.columns([4, 1])
                with cc1:
                    st.markdown(f"**{i+1}. {doc['name']}** ({doc['type']}) - {doc['size']} chars")
                with cc2:
                    if st.button(f"🗑️", key=f"del_{i}", help=f"Delete {doc['name']}"):
                        deleted_doc = st.session_state.documents.pop(i)
                        if deleted_doc['type'] == 'image' and 'image_bytes' in st.session_state:
                            if isinstance(st.session_state.image_bytes, dict):
                                st.session_state.image_bytes.pop(deleted_doc['name'], None)
                            else:
                                st.session_state.image_bytes = {}
                        rebuild_document_content()
                        st.rerun()
            
            if st.session_state.documents:
                if st.button("🔄 Reset All Documents", type="secondary"):
                    clear_all_document_state()
                    st.rerun()
        
        # ---------- Quick actions ----------
        if st.session_state.documents:
            st.markdown("**Quick Actions:**")
            q1, q2 = st.columns(2)
            with q1:
                if st.button("🗑️ Clear Chat", help="Clear chat history but keep documents"):
                    st.session_state.chat_history = []
                    st.rerun()
            with q2:
                if st.button("📊 Document Stats", help="Show dashboard with charts"):
                    st.session_state.show_stats = not st.session_state.get('show_stats', False)
                    st.rerun()

        # ---------- DASHBOARD ----------
        if st.session_state.get('show_stats', False):
            st.markdown("## 📊 Document Analytics Dashboard")

            # Data agregat per dokumen
            stats_data = []
            combined_texts = []
            struct_total = {"text":0, "tables":0, "images":0}
            for doc in st.session_state.documents:
                txt = doc['content'] or ""
                combined_texts.append(txt)
                ocr_q = _compute_ocr_quality(txt)
                words = len(_tokenize(txt))
                lines = len(txt.splitlines())
                stt = {
                    "Document": doc['name'],
                    "Type": doc['type'],
                    "Size (chars)": doc['size'],
                    "Words": words,
                    "Lines": lines,
                    "OCR Quality": round(ocr_q,1)
                }
                stats_data.append(stt)
                s = _structure_stats(txt)
                struct_total["text"] += s["text"]
                struct_total["tables"] += s["tables"]
                struct_total["images"] += s["images"]

            st.dataframe(stats_data, use_container_width=True)

            # -------- 1) Document Health --------
            st.markdown("### Document Health")
            dh1, dh2 = st.columns([1,1])
            with dh1:
                ocr_score = _compute_ocr_quality("\n\n".join(combined_texts))
                fig_g = go.Figure(go.Indicator(
                    mode="gauge+number",
                    value=ocr_score,
                    number={'suffix': " /100"},
                    title={'text': "OCR Quality Score"},
                    gauge={'axis': {'range': [0,100]},
                           'bar': {'thickness': 0.3},
                           'steps': [
                               {'range': [0,50], 'color': "#fce4ec"},
                               {'range': [50,75], 'color': "#fff3e0"},
                               {'range': [75,100], 'color': "#e8f5e9"},
                           ]}
                ))
                st.plotly_chart(fig_g, use_container_width=True)

            with dh2:
                labels = ["Text", "Tables", "Images"]
                values = [max(1, struct_total["text"]), max(1, struct_total["tables"]), max(1, struct_total["images"])]
                fig_donut = px.pie(values=values, names=labels, hole=0.6, title="Document Structure Overview")
                st.plotly_chart(fig_donut, use_container_width=True)

            secs = _extract_sections("\n\n".join(combined_texts))
            sec_labels, cat_labels, mat = _missing_matrix(secs)
            fig_heat = px.imshow(mat, aspect="auto", color_continuous_scale="Viridis",
                                 labels=dict(x="Section", y="Missing Type", color="Count"),
                                 x=sec_labels, y=cat_labels)
            fig_heat.update_layout(title="Missing Data Heatmap")
            st.plotly_chart(fig_heat, use_container_width=True)

        # ---------- Chat history ----------
        for m in st.session_state.chat_history:
            role = "You" if m["role"] == "user" else "Assistant"
            st.markdown(f"**{role}:** {m['content']}")

        # ---------- Q&A input ----------
        with st.form("qa_form_docs", clear_on_submit=True):
            user_q = st.text_input(
                "Your question (can ask about specific documents or compare them):",
                key="qa_input"
            )
            submitted = st.form_submit_button("Ask")

        if submitted and user_q:
            st.session_state.chat_history.append({"role": "user", "content": user_q})
            with st.spinner("Generating response..."):
                ans = None
                try:
                    # pilih dokumen berdasar apakah user menyebut nama file
                    docs_to_use = select_docs_for_query(user_q, st.session_state.documents)
                    avg_results = compute_average_per_doc(user_q, docs_to_use)
                except Exception:
                    avg_results = []

                if avg_results:
                    if len(avg_results) == 1:
                        r = avg_results[0]
                        val_str = f"{r['value']:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")
                        if ANSWER_LANGUAGE == "Bahasa Indonesia":
                            ans = f"Rata-rata di dokumen **{r['doc']}** untuk kolom **{r['column']}** adalah **{val_str}** (n={r['n']})."
                        else:
                            ans = f"The average in document **{r['doc']}** for column **{r['column']}** is **{val_str}** (n={r['n']})."
                    else:
                        # tampilkan semua dan bandingkan
                        ans_lines = []
                        for r in avg_results:
                            val_str = f"{r['value']:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")
                            ans_lines.append(f"- **{r['doc']}** → {val_str} (n={r['n']})")
                        best = max(avg_results, key=lambda x: x["value"])
                        if ANSWER_LANGUAGE == "Bahasa Indonesia":
                            ans = "Hasil rata-rata per dokumen:\n" + "\n".join(ans_lines)
                            ans += f"\n\n📊 Nilai tertinggi ada di **{best['doc']}** dengan rata-rata {best['value']:.2f}."
                        else:
                            ans = "Average per document:\n" + "\n".join(ans_lines)
                            ans += f"\n\n📊 Highest is in **{best['doc']}** with average {best['value']:.2f}."
                # 2) Fallback RAG jika bukan intent rata-rata / gagal parsing
                if not ans:
                    if not google_api_key:
                        ans = "Please provide a valid Google API Key."
                    else:
                        ans = generate_response_with_fallback(st.session_state.ocr_content, user_q)
            st.session_state.chat_history.append({"role": "assistant", "content": ans})
            st.rerun()  # field sudah kosong otomatis oleh form

    else:
        st.info("No documents processed yet. You can either upload files or just type a URL below and press Ask.")
        with st.form("qa_form_nodocs", clear_on_submit=True):
            user_q = st.text_input(
                "Your question (can ask about specific documents or compare them):",
                key="qa_input"
            )
            submitted = st.form_submit_button("Ask")

        if submitted and user_q:
            url_candidate = st.session_state.get("url_input")
            if url_candidate and not st.session_state.get("ocr_content"):
                with st.spinner("Processing URL before answering..."):
                    success, msg = process_url_to_content(url_candidate)
                    if not success:
                        st.error(msg)
                        st.stop()
            if st.session_state.get("ocr_content"):
                st.session_state.chat_history.append({"role": "user", "content": user_q})
                with st.spinner("Generating response..."):
                    ans = None
                    try:
                        docs_to_use = select_docs_for_query(user_q, st.session_state.documents)
                        avg_results = compute_average_per_doc(user_q, docs_to_use)
                    except Exception:
                        avg_results = []
                    if avg_results:
                        if len(avg_results) == 1:
                            r = avg_results[0]
                            val_str = f"{r['value']:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")
                            if ANSWER_LANGUAGE == "Bahasa Indonesia":
                                ans = f"Rata-rata di dokumen **{r['doc']}** untuk kolom **{r['column']}** adalah **{val_str}** (n={r['n']})."
                            else:
                                ans = f"The average in document **{r['doc']}** for column **{r['column']}** is **{val_str}** (n={r['n']})."
                        else:
                            ans_lines = []
                            for r in avg_results:
                                val_str = f"{r['value']:,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")
                                ans_lines.append(f"- **{r['doc']}** → {val_str} (n={r['n']})")
                            best = max(avg_results, key=lambda x: x["value"])
                            if ANSWER_LANGUAGE == "Bahasa Indonesia":
                                ans = "Hasil rata-rata per dokumen:\n" + "\n".join(ans_lines)
                                ans += f"\n\n📊 Nilai tertinggi ada di **{best['doc']}** dengan rata-rata {best['value']:.2f}."
                            else:
                                ans = "Average per document:\n" + "\n".join(ans_lines)
                                ans += f"\n\n📊 Highest is in **{best['doc']}** with average {best['value']:.2f}."
                    if not ans:
                        if not google_api_key:
                            ans = "Please provide a valid Google API Key."
                        else:
                            ans = generate_response_with_fallback(st.session_state.ocr_content, user_q)
                st.session_state.chat_history.append({"role": "assistant", "content": ans})
                st.rerun()  # field sudah kosong otomatis oleh form
            else:
                st.warning("Please provide a URL or upload a document first.")

# Tampilkan konten hasil OCR/ekstraksi
if st.session_state.get("documents"):
    with st.expander("📄 View All Document Contents"):
        for i, doc in enumerate(st.session_state.documents):
            st.markdown(f"### {doc['name']} ({doc['type']})")
            st.markdown(doc['content'])
            st.markdown("---")
